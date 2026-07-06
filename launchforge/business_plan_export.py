"""Business plan export helpers for the experimental Gradio UI.

The functions in this module map an existing LaunchForge launch pack into a
King's Trust-style business plan structure. They do not change launch-pack
generation; they only repackage already-generated fields into a founder-ready
document.
"""

from __future__ import annotations

import re
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List

from pydantic import BaseModel

from launchforge.config import as_money
from launchforge.schemas import model_to_dict


def _to_plain(value: Any) -> Any:
    if isinstance(value, BaseModel):
        return model_to_dict(value)
    if isinstance(value, dict):
        return {key: _to_plain(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_to_plain(item) for item in value]
    return value


def _items(values: Any, limit: int | None = None) -> List[str]:
    if not values:
        return []
    if isinstance(values, str):
        output = [values]
    else:
        output = [str(item) for item in values if str(item).strip()]
    return output[:limit] if limit else output


def _money(data: Dict[str, Any], amount: Any) -> str:
    try:
        value = float(amount or 0)
    except (TypeError, ValueError):
        value = 0.0
    return as_money(value, data.get("currency_symbol", "\u00a3"))


def _business_label(business_type: str) -> str:
    labels = {
        "local_service": "Local service",
        "physical_retail": "Physical retail",
        "ecommerce": "E-commerce",
        "digital_product": "Digital product",
        "food_drink": "Food and drink",
        "b2b_service": "B2B service",
        "event_community": "Event or community",
        "unknown": "Business model to validate",
    }
    return labels.get(business_type or "unknown", str(business_type or "Business").replace("_", " ").title())


def _input_dict(data: Dict[str, Any], founder_input: Any = None) -> Dict[str, Any]:
    source = _to_plain(founder_input) if founder_input else data.get("input", {})
    return source if isinstance(source, dict) else {}


def _business_name(data: Dict[str, Any]) -> str:
    idea = str((data.get("input") or {}).get("idea") or "").strip()
    if not idea:
        return "LaunchForge Draft Business"
    lowered = idea.lower()
    if "tutor" in lowered:
        return "Structured STEM Tutoring"
    if "corner shop" in lowered or "shop near" in lowered:
        return "Station Essentials"
    if "shopify" in lowered or "store" in lowered:
        return "ForgeFit Accessories"
    words = re.findall(r"[A-Za-z0-9]+", idea)
    meaningful = [word for word in words if word.lower() not in {"i", "want", "to", "start", "launch", "open", "a", "an", "the", "business"}]
    return " ".join(meaningful[:3]).title() or "LaunchForge Draft Business"


def _first_persona(data: Dict[str, Any]) -> Dict[str, Any]:
    scores = data.get("segment_scores") or []
    recommended = next((item for item in scores if item.get("recommended_first_segment")), None)
    personas = data.get("personas") or []
    if recommended:
        persona = next((item for item in personas if item.get("name") == recommended.get("persona_name")), {})
        return {**persona, **recommended}
    return personas[0] if personas else {}


def _first_offer(data: Dict[str, Any]) -> Dict[str, Any]:
    offers = data.get("offer_ladder") or []
    return offers[0] if offers else {}


def _strapline(data: Dict[str, Any]) -> str:
    hooks = (data.get("marketing_messages") or {}).get("hooks") or []
    if hooks:
        return str(hooks[0]).rstrip(".")
    segment = _first_persona(data).get("segment") or "your first customers"
    return f"A focused launch plan for {segment}"


def _elevator_pitch(data: Dict[str, Any]) -> str:
    classification = data.get("classification") or {}
    persona = _first_persona(data)
    offer = _first_offer(data)
    business_type = _business_label(classification.get("business_type", "unknown")).lower()
    target = persona.get("segment") or (data.get("input") or {}).get("target_customer") or "a defined customer segment"
    pain = "; ".join(_items(persona.get("pains"), 2)) or "a clear problem they need solved"
    outcome = offer.get("success_metric") or "a measurable next step"
    offer_name = offer.get("name") or "first offer"
    return (
        f"{_business_name(data)} is a {business_type} helping {target} address {pain}. "
        f"The first offer is {offer_name}, designed to create {outcome} while validating demand before scaling."
    )


def _markdown_table(headers: List[str], rows: Iterable[Iterable[Any]]) -> str:
    header = "| " + " | ".join(headers) + " |"
    divider = "| " + " | ".join("---" for _ in headers) + " |"
    body = ["| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |" for row in rows]
    return "\n".join([header, divider, *body])


def _forecast_12_months(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    cashflow = list(data.get("cashflow") or [])
    rows: List[Dict[str, Any]] = []
    previous_closing = 0.0
    if cashflow:
        first = cashflow[0]
        previous_closing = float(first.get("cumulative_cashflow", 0)) - float(first.get("net_cashflow", 0))
    for month in range(1, 13):
        if month <= len(cashflow):
            source = cashflow[month - 1]
            revenue = float(source.get("revenue", 0))
            costs = float(source.get("costs", 0))
            net = float(source.get("net_cashflow", revenue - costs))
            closing = float(source.get("cumulative_cashflow", previous_closing + net))
        else:
            previous = rows[-1] if rows else {"revenue": 0.0, "costs": 0.0, "closing_balance": 0.0}
            revenue = float(previous["revenue"]) * 1.05
            costs = float(previous["costs"]) * 1.03
            net = revenue - costs
            closing = float(previous["closing_balance"]) + net
        rows.append(
            {
                "month": month,
                "revenue": round(revenue, 2),
                "costs": round(costs, 2),
                "net_cashflow": round(net, 2),
                "opening_balance": round(previous_closing, 2),
                "closing_balance": round(closing, 2),
            }
        )
        previous_closing = closing
    return rows


def _cost_rows(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    costs = data.get("startup_costs") or {}
    if costs:
        return [{"item": key, "amount": value, "notes": "Launch/startup planning cost"} for key, value in costs.items()]
    return [
        {"item": "Materials/equipment", "amount": 0, "notes": "To validate"},
        {"item": "Marketing", "amount": 0, "notes": "To validate"},
        {"item": "Software/admin", "amount": 0, "notes": "To validate"},
        {"item": "Insurance/legal", "amount": 0, "notes": "To validate"},
        {"item": "Contingency", "amount": 0, "notes": "To validate"},
    ]


def _swot(data: Dict[str, Any]) -> Dict[str, List[str]]:
    strengths = _items(data.get("readiness_strengths"), 4) or _items((data.get("classification") or {}).get("matched_signals"), 4)
    gaps = _items(data.get("readiness_gaps"), 4)
    risks = _items(data.get("risks"), 4)
    segment = _first_persona(data)
    opportunities = [
        f"Focus first on {segment.get('segment', 'the highest-fit customer segment')}.",
        "Use validation conversations and early proof assets to reduce launch uncertainty.",
    ]
    return {
        "Strengths": strengths or ["Clear launch concept and initial business model route."],
        "Weaknesses": gaps or ["Founder must complete validation and operating setup."],
        "Opportunities": opportunities,
        "Threats": risks or ["Competitors, costs, and customer conversion still need validation."],
    }


def build_business_plan_dict(pack: Any, founder_input: Any = None) -> Dict[str, Any]:
    """Return a structured business plan dictionary from a launch pack."""

    data = _to_plain(pack)
    if not isinstance(data, dict):
        data = {}
    founder = _input_dict(data, founder_input)
    classification = data.get("classification") or {}
    persona = _first_persona(data)
    offer = _first_offer(data)
    business_type = classification.get("business_type", "unknown")
    forecast = _forecast_12_months(data)
    startup_total = sum(float(value or 0) for value in (data.get("startup_costs") or {}).values())
    pricing = data.get("pricing") or []
    marketing = data.get("marketing_messages") or {}
    funnel = data.get("funnel_model") or data.get("sales_funnel", {}).get("stages") or []
    swot = _swot(data)
    competitor_rows = [
        {"competitor": "To research: local/direct competitors", "offer": "To validate", "price": "To validate", "advantage": "Compare proof, convenience, and pricing"},
        {"competitor": "To research: online/substitute competitors", "offer": "To validate", "price": "To validate", "advantage": "Compare reach, speed, and customer trust"},
    ]
    plan = {
        "business_name": _business_name(data),
        "owner_name": "To be completed",
        "business_address": "To be completed",
        "business_email": "To be completed",
        "business_phone": "To be completed",
        "strapline": _strapline(data),
        "elevator_pitch": _elevator_pitch(data),
        "business_type": _business_label(business_type),
        "currency_code": data.get("currency_code", "GBP"),
        "currency_symbol": data.get("currency_symbol", "\u00a3"),
        "sections": {
            "getting_started": {
                "Business name": _business_name(data),
                "Owner name": "To be completed",
                "Business address": "To be completed",
                "Business email": "To be completed",
                "Business phone": "To be completed",
            },
            "1_executive_summary": {
                "1.1 Business summary": (
                    f"{founder.get('idea') or (data.get('input') or {}).get('idea') or 'To be completed by founder'} "
                    f"The model is classified as {_business_label(business_type)} with "
                    f"{int(float(classification.get('confidence', 0)) * 100)}% confidence. "
                    f"The first customer focus is {persona.get('segment', 'to be validated')} and the first offer is "
                    f"{offer.get('name', 'to be completed')}."
                ),
                "1.2 Business aims": _items(data.get("next_3_actions"), 3)
                + [f"Complete the 30-day roadmap covering {len(data.get('roadmap') or [])} launch tasks."],
                "1.3 Financial summary": [
                    f"Estimated startup cost: {_money(data, startup_total)}.",
                    f"Modelled break-even month: {data.get('breakeven_month', 'To validate')}.",
                    f"Readiness score: {data.get('readiness_score', 'To validate')}/100.",
                    data.get("forecast_disclaimer", "Planning forecast only; not financial advice."),
                ],
            },
            "elevator_pitch": {
                "1.4 Business name": _business_name(data),
                "1.5 Strapline": _strapline(data),
                "1.6 Elevator pitch": _elevator_pitch(data),
            },
            "2_owner_background": {
                "Why run the business": "To be completed by founder",
                "Work experience": founder.get("founder_resources") or "To be completed by founder",
                "Qualifications and education": "To be completed by founder",
                "Training": "To be completed by founder",
                "Hobbies/interests": "To be completed by founder",
                "Additional information": "No personal qualifications have been invented by LaunchForge.",
            },
            "3_products_and_services": data.get("offer_ladder") or [],
            "4_the_market": {
                "target_customer": founder.get("target_customer") or persona.get("segment") or "To validate",
                "personas": data.get("personas") or [],
                "segment_scores": data.get("segment_scores") or [],
                "location": founder.get("location") or (data.get("input") or {}).get("location") or "To validate",
            },
            "5_market_research": {
                "classification_evidence": _items(classification.get("matched_signals"), 8),
                "assumptions": _items(data.get("assumptions"), 8),
                "validation_gaps": _items(data.get("readiness_gaps"), 8)
                or ["Field research, competitor checks, and test sales still need to be completed."],
            },
            "6_marketing_strategy": {
                "channels": (data.get("sales_funnel") or {}).get("channels") or [],
                "funnel": funnel,
                "hooks": marketing.get("hooks") or [],
                "social_posts": marketing.get("social_posts") or [],
                "whatsapp_email": marketing.get("whatsapp_email") or [],
            },
            "7_competitor_analysis": {
                "competitors": competitor_rows,
                "swot": swot,
                "usp": offer.get("description") or "The USP should be sharpened after competitor research.",
            },
            "8_operations_and_logistics": {
                "operations_checklist": data.get("operations_checklist") or [],
                "capacity_model": data.get("capacity_model") or {},
                "payment_methods": _payment_methods(business_type),
                "suppliers_equipment_legal": "Suppliers, equipment, insurance, and legal requirements to be confirmed by founder.",
            },
            "9_costs_and_pricing_strategy": {
                "pricing": pricing,
                "pricing_scenarios": data.get("pricing_scenarios") or [],
                "startup_costs": _cost_rows(data),
                "assumptions": (data.get("cashflow_assumptions") or {}).get("costs", []),
            },
            "10_financial_forecasts": {
                "10.1 Sales and costs forecast": forecast,
                "10.2 Personal survival budget": _personal_survival_budget_rows(),
                "10.3 Cashflow forecast": forecast,
                "10.4 Costs table": _cost_rows(data),
                "disclaimer": data.get("forecast_disclaimer", "Planning model only; not financial advice."),
            },
            "11_backup_plan": {
                "short_term": _items(data.get("risks"), 3) or ["Validate demand before committing more spend."],
                "long_term": _items([task.get("outcome") for task in data.get("roadmap", []) if isinstance(task, dict)], 3),
                "plan_b": _items(data.get("critic_notes"), 3)
                or ["Pause discretionary spend, return to customer interviews, and simplify the first offer."],
            },
        },
    }
    return plan


def _payment_methods(business_type: str) -> List[str]:
    if business_type == "physical_retail":
        return ["Card terminal", "Cash", "Mobile wallet"]
    if business_type == "ecommerce":
        return ["Shopify Payments", "PayPal", "Card checkout"]
    return ["Bank transfer", "Card payment link", "Booking checkout"]


def _personal_survival_budget_rows() -> List[Dict[str, str]]:
    return [
        {"category": "Rent/mortgage", "amount": "To be completed by founder"},
        {"category": "Utilities", "amount": "To be completed by founder"},
        {"category": "Food and household", "amount": "To be completed by founder"},
        {"category": "Transport", "amount": "To be completed by founder"},
        {"category": "Debt repayments", "amount": "To be completed by founder"},
        {"category": "Tax/benefits/other", "amount": "To be completed by founder"},
    ]


def _bullets(values: Any) -> str:
    items = _items(values)
    if not items:
        return "- To be completed"
    return "\n".join(f"- {item}" for item in items)


def generate_business_plan_markdown(pack: Any, founder_input: Any = None) -> str:
    """Generate a Markdown business plan using the King's Trust-style structure."""

    plan = build_business_plan_dict(pack, founder_input)
    sections = plan["sections"]
    data = _to_plain(pack)
    if not isinstance(data, dict):
        data = {}
    lines = [
        f"# {plan['business_name']} Business Plan",
        "",
        "Business plan generated using the King's Trust-style section structure.",
        "",
        "_Financial figures are planning assumptions only and are not financial advice._",
        "",
        "## Getting started / business and owner details",
    ]
    for label, value in sections["getting_started"].items():
        lines.append(f"- **{label}:** {value}")

    executive = sections["1_executive_summary"]
    lines.extend(["", "## Section 1: Executive summary", "", "### 1.1 Business summary", executive["1.1 Business summary"], ""])
    lines.extend(["### 1.2 Business aims", _bullets(executive["1.2 Business aims"]), "", "### 1.3 Financial summary", _bullets(executive["1.3 Financial summary"])])

    pitch = sections["elevator_pitch"]
    lines.extend(
        [
            "",
            "## Elevator pitch",
            "",
            f"### 1.4 Business name\n{pitch['1.4 Business name']}",
            "",
            f"### 1.5 Strapline\n{pitch['1.5 Strapline']}",
            "",
            f"### 1.6 Elevator pitch\n{pitch['1.6 Elevator pitch']}",
        ]
    )

    owner = sections["2_owner_background"]
    lines.extend(["", "## Section 2: Owner's background"])
    for label, value in owner.items():
        lines.extend(["", f"### {label.replace('_', ' ').title()}", str(value)])

    lines.extend(["", "## Section 3: Products and services"])
    offers = sections["3_products_and_services"]
    if offers:
        for offer in offers:
            lines.extend(
                [
                    "",
                    f"### {offer.get('name', 'Offer')}",
                    offer.get("description", "To be completed"),
                    "",
                    f"**Ideal for:** {offer.get('ideal_for', 'To validate')}",
                    "",
                    "**Deliverables:**",
                    _bullets(offer.get("deliverables")),
                    "",
                    f"**Success metric:** {offer.get('success_metric', 'To validate')}",
                ]
            )
    else:
        lines.append("To be completed by founder.")

    market = sections["4_the_market"]
    lines.extend(["", "## Section 4: The market"])
    lines.append(f"Primary target customer: {market.get('target_customer', 'To validate')}.")
    lines.append(f"Location/channel context: {market.get('location', 'To validate')}.")
    persona_rows = [
        [
            item.get("name", "Persona"),
            item.get("segment", "To validate"),
            "; ".join(_items(item.get("pains"), 2)),
            "; ".join(_items(item.get("channels"), 3)),
            item.get("buying_trigger", "To validate"),
        ]
        for item in market.get("personas", [])
    ]
    if persona_rows:
        lines.extend(["", _markdown_table(["Persona", "Segment", "Pain", "Channels", "Trigger"], persona_rows)])
    score_rows = [
        [
            item.get("persona_name", item.get("segment", "Segment")),
            item.get("pain_intensity", "-"),
            item.get("reachability", "-"),
            item.get("urgency", "-"),
            item.get("willingness_to_pay", "-"),
            item.get("overall_score", "-"),
        ]
        for item in market.get("segment_scores", [])
    ]
    if score_rows:
        lines.extend(["", _markdown_table(["Segment", "Pain", "Reach", "Urgency", "Pay", "Overall"], score_rows)])

    research = sections["5_market_research"]
    lines.extend(["", "## Section 5: Market research", "", "### Classification evidence", _bullets(research["classification_evidence"]), "", "### Assumptions", _bullets(research["assumptions"]), "", "### Field research to complete", _bullets(research["validation_gaps"])])

    strategy = sections["6_marketing_strategy"]
    channel_rows = [[channel, "To validate", "To validate"] for channel in strategy.get("channels", [])]
    if not channel_rows:
        channel_rows = [["To validate", "To validate", "To validate"]]
    lines.extend(["", "## Section 6: Marketing strategy", _markdown_table(["Channel", "Message / test", "Budget"], channel_rows), "", "### Hooks", _bullets(strategy.get("hooks")), "", "### Social posts", _bullets(strategy.get("social_posts")), "", "### WhatsApp / Email", _bullets(strategy.get("whatsapp_email"))])

    competitors = sections["7_competitor_analysis"]
    lines.extend(
        [
            "",
            "## Section 7: Competitor analysis",
            _markdown_table(
                ["Competitor", "Offer", "Price", "What to compare"],
                [[row["competitor"], row["offer"], row["price"], row["advantage"]] for row in competitors["competitors"]],
            ),
            "",
            "### SWOT",
        ]
    )
    for label, values in competitors["swot"].items():
        lines.extend(["", f"**{label}**", _bullets(values)])
    lines.extend(["", "### USP", competitors["usp"]])

    operations = sections["8_operations_and_logistics"]
    capacity = operations.get("capacity_model") or {}
    lines.extend(["", "## Section 8: Operations and logistics", "", "### Operating checklist", _bullets(operations.get("operations_checklist")), "", "### Capacity model", _markdown_table(["Item", "Value"], [[key.replace("_", " ").title(), value] for key, value in capacity.items()] or [["To complete", "To validate"]]), "", "### Payment methods", _bullets(operations.get("payment_methods")), "", operations["suppliers_equipment_legal"]])

    pricing_section = sections["9_costs_and_pricing_strategy"]
    pricing_rows = [
        [
            item.get("tier", "Tier"),
            _money(data, item.get("price", 0)),
            item.get("unit", ""),
            item.get("rationale", ""),
            item.get("upgrade_path", ""),
        ]
        for item in pricing_section.get("pricing", [])
    ]
    lines.extend(["", "## Section 9: Costs and pricing strategy"])
    if pricing_rows:
        lines.append(_markdown_table(["Tier", "Price", "Unit", "Rationale", "Upgrade path"], pricing_rows))
    else:
        lines.append("Pricing to be completed by founder.")
    cost_rows = [[row["item"], _money(data, row["amount"]), row["notes"]] for row in pricing_section.get("startup_costs", [])]
    lines.extend(["", "### Startup costs", _markdown_table(["Item", "Amount", "Notes"], cost_rows), "", "### Cost assumptions", _bullets(pricing_section.get("assumptions"))])

    forecasts = sections["10_financial_forecasts"]
    forecast_rows = [
        [
            row["month"],
            _money(data, row["revenue"]),
            _money(data, row["costs"]),
            _money(data, row["net_cashflow"]),
            _money(data, row["opening_balance"]),
            _money(data, row["closing_balance"]),
        ]
        for row in forecasts["10.1 Sales and costs forecast"]
    ]
    lines.extend(
        [
            "",
            "## Section 10: Financial forecasts",
            "",
            "### 10.1 Sales and costs forecast",
            _markdown_table(["Month", "Sales", "Costs", "Net", "Opening balance", "Closing balance"], forecast_rows),
            "",
            "### 10.2 Personal survival budget",
            _markdown_table(["Category", "Amount"], [[row["category"], row["amount"]] for row in forecasts["10.2 Personal survival budget"]]),
            "",
            "### 10.3 Cashflow forecast",
            _markdown_table(["Month", "Revenue", "Costs", "Net cashflow", "Closing balance"], [[row["month"], _money(data, row["revenue"]), _money(data, row["costs"]), _money(data, row["net_cashflow"]), _money(data, row["closing_balance"])] for row in forecasts["10.3 Cashflow forecast"]]),
            "",
            "### 10.4 Costs table",
            _markdown_table(["Item", "Amount", "Notes"], [[row["item"], _money(data, row["amount"]), row["notes"]] for row in forecasts["10.4 Costs table"]]),
            "",
            f"_{forecasts['disclaimer']} Planning model only._",
        ]
    )

    backup = sections["11_backup_plan"]
    lines.extend(["", "## Section 11: Back-up Plan", "", "### Short-term risk response", _bullets(backup["short_term"]), "", "### Longer-term adjustment", _bullets(backup["long_term"]), "", "### Plan B", _bullets(backup["plan_b"])])
    return "\n".join(lines).strip() + "\n"


def generate_business_plan_docx(pack: Any, founder_input: Any = None, output_path: str | None = None) -> str:
    """Generate a DOCX business plan.

    Raises RuntimeError when python-docx is not installed. The Gradio UI catches
    that error and provides a Markdown fallback.
    """

    try:
        from docx import Document
        from docx.shared import Pt
    except ModuleNotFoundError as exc:  # pragma: no cover - depends on optional package
        raise RuntimeError("python-docx is required for DOCX export. Use the Markdown business plan export or install python-docx.") from exc

    markdown = generate_business_plan_markdown(pack, founder_input)
    if output_path:
        path = Path(output_path)
    else:
        export_dir = Path(tempfile.gettempdir()) / "launchforge_exports"
        export_dir.mkdir(parents=True, exist_ok=True)
        stamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S_%f")
        path = export_dir / f"launchforge_business_plan_{stamp}.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("| ") and " |" in line:
            continue
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("_") and line.endswith("_"):
            document.add_paragraph(line.strip("_"))
        else:
            document.add_paragraph(line)
    _add_docx_tables(document, build_business_plan_dict(pack, founder_input))
    document.save(path)
    if not path.exists():
        raise OSError(f"DOCX export was not created: {path}")
    return str(path)


def _add_docx_tables(document: Any, plan: Dict[str, Any]) -> None:  # pragma: no cover - exercised when docx is installed
    sections = plan["sections"]
    document.add_page_break()
    document.add_heading("Structured Tables", level=1)

    def add_table(title: str, headers: List[str], rows: List[List[Any]]) -> None:
        document.add_heading(title, level=2)
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = str(header)
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = str(value)

    strategy = sections["6_marketing_strategy"]
    add_table("Marketing strategy", ["Channel", "Message / test", "Budget"], [[channel, "To validate", "To validate"] for channel in strategy.get("channels", [])] or [["To validate", "To validate", "To validate"]])

    competitors = sections["7_competitor_analysis"]["competitors"]
    add_table("Competitor analysis", ["Competitor", "Offer", "Price", "What to compare"], [[row["competitor"], row["offer"], row["price"], row["advantage"]] for row in competitors])

    swot = sections["7_competitor_analysis"]["swot"]
    add_table("SWOT", ["Strengths", "Weaknesses", "Opportunities", "Threats"], [[
        "\n".join(swot.get("Strengths", [])),
        "\n".join(swot.get("Weaknesses", [])),
        "\n".join(swot.get("Opportunities", [])),
        "\n".join(swot.get("Threats", [])),
    ]])

    costs = sections["10_financial_forecasts"]["10.4 Costs table"]
    add_table("Costs table", ["Item", "Amount", "Notes"], [[row["item"], row["amount"], row["notes"]] for row in costs])

    forecast = sections["10_financial_forecasts"]["10.1 Sales and costs forecast"]
    add_table("Sales and costs forecast", ["Month", "Revenue", "Costs", "Net", "Closing"], [[row["month"], row["revenue"], row["costs"], row["net_cashflow"], row["closing_balance"]] for row in forecast])
